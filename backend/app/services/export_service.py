import io
from abc import ABC, abstractmethod

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.project import Project
from app.models.chapter import Chapter
from app.models.character import Character


class BaseExporter(ABC):
    @abstractmethod
    async def export(
        self, db: AsyncSession, project_id: str, options: dict | None = None
    ) -> bytes: ...


class TxtExporter(BaseExporter):
    async def export(
        self, db: AsyncSession, project_id: str, options: dict | None = None
    ) -> bytes:
        project = await db.get(Project, project_id)
        chapters = (
            await db.execute(
                select(Chapter)
                .where(Chapter.project_id == project_id)
                .order_by(Chapter.sort_order)
            )
        ).scalars().all()
        buf = io.StringIO()
        buf.write(f"{project.name}\n{'=' * 40}\n\n")
        if options and options.get("include_outline"):
            buf.write("[大纲]\n\n")
        if options and options.get("include_characters"):
            characters = (
                await db.execute(
                    select(Character).where(
                        Character.project_id == project_id
                    ).order_by(
                        Character.sort_order, Character.created_at, Character.id
                    )
                )
            ).scalars().all()
            buf.write("[人物档案]\n")
            for c in characters:
                buf.write(f"  {c.name}\n")
                if c.biography:
                    buf.write(f"    人物小传：{c.biography}\n")
            buf.write("\n")
        for ch in chapters:
            buf.write(
                f"第{ch.sort_order + 1}章 {ch.title}\n{'-' * 30}\n\n{ch.content or ''}\n\n"
            )
        return buf.getvalue().encode("utf-8")


class MarkdownExporter(BaseExporter):
    async def export(
        self, db: AsyncSession, project_id: str, options: dict | None = None
    ) -> bytes:
        project = await db.get(Project, project_id)
        chapters = (
            await db.execute(
                select(Chapter)
                .where(Chapter.project_id == project_id)
                .order_by(Chapter.sort_order)
            )
        ).scalars().all()
        buf = io.StringIO()
        buf.write(f"# {project.name}\n\n> {project.description or ''}\n\n")
        if options and options.get("include_characters"):
            characters = (
                await db.execute(
                    select(Character).where(
                        Character.project_id == project_id
                    ).order_by(
                        Character.sort_order, Character.created_at, Character.id
                    )
                )
            ).scalars().all()
            buf.write("## 人物档案\n\n")
            for c in characters:
                buf.write(f"- **{c.name}**\n")
                if c.biography:
                    buf.write(f"  - 人物小传：{c.biography}\n")
            buf.write("\n")
        for ch in chapters:
            buf.write(
                f"## 第{ch.sort_order + 1}章 {ch.title}\n\n{ch.content or ''}\n\n"
            )
        return buf.getvalue().encode("utf-8")


class DocxExporter(BaseExporter):
    async def export(
        self, db: AsyncSession, project_id: str, options: dict | None = None
    ) -> bytes:
        from docx import Document

        project = await db.get(Project, project_id)
        chapters = (
            await db.execute(
                select(Chapter)
                .where(Chapter.project_id == project_id)
                .order_by(Chapter.sort_order)
            )
        ).scalars().all()
        doc = Document()
        doc.add_heading(project.name, level=0)
        for ch in chapters:
            doc.add_heading(
                f"第{ch.sort_order + 1}章 {ch.title}", level=1
            )
            for para in (ch.content or "").split("\n"):
                if para.strip():
                    doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


EXPORTER_MAP = {
    "txt": TxtExporter,
    "markdown": MarkdownExporter,
    "docx": DocxExporter,
}


class ExportService:
    async def export_project(
        self,
        db: AsyncSession,
        project_id: str,
        format: str,
        options: dict | None = None,
    ) -> tuple[bytes, str, str]:
        exporter_cls = EXPORTER_MAP.get(format, TxtExporter)
        exporter = exporter_cls()
        content = await exporter.export(db, project_id, options)
        project = await db.get(Project, project_id)

        content_type_map = {
            "txt": "text/plain; charset=utf-8",
            "markdown": "text/markdown; charset=utf-8",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf",
        }
        ext_map = {"txt": ".txt", "markdown": ".md", "docx": ".docx", "pdf": ".pdf"}

        filename = f"{project.name}{ext_map.get(format, '.txt')}"
        content_type = content_type_map.get(format, "text/plain; charset=utf-8")
        return content, filename, content_type


export_service = ExportService()
